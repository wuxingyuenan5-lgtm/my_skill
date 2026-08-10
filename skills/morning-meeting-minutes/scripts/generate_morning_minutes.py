#!/usr/bin/env python3
"""生成 2026年8月7日会议纪要（一页纸，紧凑合并版）"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_TITLE = RGBColor(0x1F, 0x38, 0x64)
COLOR_SUB = RGBColor(0x1F, 0x38, 0x64)
COLOR_TEXT = RGBColor(0x00, 0x00, 0x00)
FONT_CN = '宋体'  # 跟用户截图的宋体字风格一致
FONT_EN = 'Times New Roman'


def set_run_font(run, size_pt, color=None, bold=False):
    run.font.size = Pt(size_pt)
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    b_el = rPr.find(qn('w:b'))
    if bold:
        if b_el is None:
            rPr.append(OxmlElement('w:b'))
    else:
        if b_el is not None:
            rPr.remove(b_el)


def set_para_spacing(p, before=2, after=2, line=1.25, indent_cm=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)


def add_paragraph(doc, text, *, size=10, color=COLOR_TEXT, bold=False,
                  align=None, indent_cm=None, before=2, after=2, line=1.25):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, before=before, after=after, line=line, indent_cm=indent_cm)
    run = p.add_run(text)
    set_run_font(run, size, color, bold)
    return p


def add_centered_title(doc, text, size=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=6, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size, COLOR_TITLE, bold=True)
    return p


def add_section_heading(doc, text, size=12):
    return add_paragraph(doc, text, size=size, color=COLOR_SUB, bold=True,
                         before=8, after=3, line=1.2)


def add_sub_heading(doc, text, size=11):
    return add_paragraph(doc, text, size=size, color=COLOR_SUB, bold=True,
                         before=5, after=2, line=1.2)


def add_para_block(doc, parts, *, indent_cm=0.3, before=1, after=1):
    """parts: list of (text, bold) 元组，拼成一段"""
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after, line=1.3, indent_cm=indent_cm)
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, 10, COLOR_TEXT, bold)
    return p


def set_default_font(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    style.font.size = Pt(10)


def main():
    doc = Document()
    set_default_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # 单标题
    add_centered_title(doc, '2026.08.07 会议纪要', size=15)

    # 一、大盘 / 夜盘 / 商品
    add_section_heading(doc, '一、大盘 / 夜盘 / 商品')
    add_paragraph(doc, '海外夜盘美股表现平稳，金属温和回落。', indent_cm=0.3, after=3)

    # 二、核心讨论
    add_section_heading(doc, '二、核心讨论')

    add_sub_heading(doc, '2.1 法案进展与板块差异')
    add_para_block(doc, [
        ('FCC 法案有进展，关键看背后逻辑数据。', False),
        ('光模块：', True),
        ('符合法案，受负面影响；', False),
        ('PCB：', True),
        ('不符合法案，风险不大，存在加仓机会，', False),
        ('胜宏科技（PCB 龙头）可修复到 5000 亿估值。', True),
    ])
    add_para_block(doc, [
        ('产业链关注：', True),
        ('蓝亚（估值较贵）、芯碁微装。', False),
    ])

    add_sub_heading(doc, '2.2 资源民族主义主线')
    add_para_block(doc, [
        ('宁德时代相关法案使资源板块更稳，资源民族主义逻辑强化；煤炭上涨与之相关，', False),
        ('洛阳钼业可考虑配置。', True),
    ])

    add_sub_heading(doc, '2.3 公司基金经理观点与操作策略')
    add_para_block(doc, [
        ('券商：', True),
        ('偏成长型、性价比不高，PE 上升空间有限；fof 推荐并配置建滔。', False),
    ])
    add_para_block(doc, [
        ('电力：', True),
        ('随 AI 产业链上涨被提及，作为大票修复方向之一。', False),
    ])
    add_para_block(doc, [
        ('医药：', True),
        ('走势有点难看但持仓尚可；公司未特别看好，逻辑不顺。', False),
    ])
    add_para_block(doc, [
        ('操作原则：', True),
        ('跟踪哪边强就重仓哪边，不预测，跟踪验证。', False),
    ])

    # 三、观察的 fof 子基金
    add_section_heading(doc, '三、观察的 fof 子基金（可吸取优点）')
    add_para_block(doc, [
        ('杠杆 + 择时策略：', True),
        ('平均杠杆 120–130%，最高 150%，在范围内做择时；持仓周期平均 30–40 个交易日，节奏稳定。', False),
    ])
    add_para_block(doc, [
        ('交易层面极强：', True),
        ('不依赖分析，靠交易执行力；风控是命门；被套不轻易止损，仓位逐步慢加。', False),
    ])
    add_para_block(doc, [
        ('板块轮动思维：', True),
        ('无选择偏好 + 严格筛选标准 + 仓位分配。', False),
    ])
    add_para_block(doc, [
        ('值得借鉴：', True),
        ('交易纪律、波动控制、回撤意识、择时能力。', False),
    ])

    # 四、后续任务
    add_section_heading(doc, '四、后续任务')
    add_para_block(doc, [
        ('自选股按行业分类，梳理核心题材，标注股价上涨催化剂、一致性预期和市值空间；', False),
        ('对 160+ 只咨询股按行业做分类，关注景气度跟踪而非单纯游资催化。', False),
    ])

    output_path = '/Users/zhangxu/WorkBuddy/2026-08-05-09-12-17/2026年8月7日晨会纪要.docx'
    doc.save(output_path)
    print('saved:', output_path)


if __name__ == '__main__':
    main()
